import re
import language_tool_python
from copy import deepcopy
from docx import Document
from docx.shared import RGBColor, Pt

# Define light blue color for emails (RGB for light blue #ADD8E6)
LIGHT_BLUE = RGBColor(173, 216, 230)

class DeterministicPreprocessor:
    def __init__(self):
        self.style_guide = {
            "paragraph": {
                "font_size": 12,
                "font_colors": ["000000", "0000FF"],
                "font_name": ["Calibri", "Calibri Light"],
                "bold": True,
                "italic": False,
            },
            "spacing": {
                "no_space_before_symbol": True,
                "paragraph_spacing_before": 6,
                "paragraph_spacing_after": 6,
                "line_spacing": 1.5,
            },
            "emails": {"font_color": "ADD8E6", "underline": True},
            "acronyms": {"expand_on_first_use": True},
            "spelling": {"check_spelling": True},
            "abbreviations": {"check_rank_abbreviations": True},
            "corrections": {"fix_automatically": True, "log_changes": True},
        }

        # Initialize LanguageTool for English
        try:
            self.tool = language_tool_python.LanguageTool("en-AU")
        except Exception as e:
            raise RuntimeError(f"Error initializing LanguageTool: {e}")

    def get_paragraph_font_info(self, paragraph):
        try:
            if paragraph.runs:
                para_style = paragraph.style
                run = paragraph.runs[0]
                font_name = run.font.name if run.font.name else para_style.font.name
                font_size = (
                    run.font.size.pt
                    if run.font.size
                    else para_style.font.size.pt if para_style.font.size else None
                )
                font_color = (
                    run.font.color.rgb
                    if run.font.color and run.font.color.rgb
                    else "Default"
                )
            else:
                font_name = paragraph.style.font.name
                font_size = paragraph.style.font.size.pt if paragraph.style.font.size else None
                font_color = "Default"
        except Exception as e:
            raise RuntimeError(f"Error reading paragraph font information: {e}")

        return {
            "font_name": font_name if font_name else "Default",
            "font_size": font_size if font_size else "Default",
            "font_color": font_color,
        }

    def detect_abbreviations(self, paragraph):
        corrections = []
        try:
            text = paragraph.text
            abbrev_pattern = r"\b[A-Z]{2,}\b"
            abbreviations = re.findall(abbrev_pattern, text)

            for abbr in abbreviations:
                matches = self.tool.check(abbr)
                if not matches:
                    corrections.append(
                        {
                            "before": abbr,
                            "after": "Requires full form on first occurrence",
                            "text": paragraph.text,
                            "issue": f"Abbreviation '{abbr}' detected. Ensure full form is used on first occurrence.",
                        }
                    )
        except Exception as e:
            raise RuntimeError(f"Error detecting abbreviations: {e}")

        return corrections

    def correct_spelling_and_grammar(self, paragraph, proper_nouns):
        corrections = []
        try:
            text = paragraph.text
            matches = self.tool.check(text)

            # Use LanguageTool to correct the text
            corrected_text = self.tool.correct(text)
            # print(f"text: {text}, -- corrected text :{corrected_text}")

            # Only proceed if corrections were made
            if corrected_text != text:
                corrections.append(
                    {
                        "before": text,
                        "after": corrected_text,
                        "text": paragraph.text,
                        "issue": "Spelling and grammar issues corrected."
                    }
                )

                # Update the paragraph text with the corrected version
                paragraph.clear()  # Clear current content
                paragraph.add_run(corrected_text)  # Add corrected content

        except Exception as e:
            raise RuntimeError(f"Error checking spelling and grammar: {e}")

        return corrections

    def detect_proper_nouns(self, paragraph):
        proper_nouns = set()
        try:
            words = paragraph.text.split()
            for word in words:
                if word and word[0].isupper():
                    proper_nouns.add(word)
        except Exception as e:
            raise RuntimeError(f"Error detecting proper nouns: {e}")

        return proper_nouns

    def format_emails(self, paragraph):
        corrections = []
        try:
            email_pattern = r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"

            for run in paragraph.runs:
                if re.search(email_pattern, run.text):
                    corrections.append(
                        {
                            "before": run.text,
                            "after": run.text,
                            "text": paragraph.text,
                            "issue": f"Email formatting: '{run.text}' has been updated with light blue color and underline.",
                        }
                    )
                    run.font.color.rgb = LIGHT_BLUE
                    run.font.underline = True
        except Exception as e:
            raise RuntimeError(f"Error formatting emails: {e}")

        return corrections

    def correct_font_size(self, paragraph):
        corrections = []
        try:
            para_info = self.get_paragraph_font_info(paragraph)
            expected_font_size = self.style_guide["paragraph"]["font_size"]

            if (
                para_info["font_size"] != expected_font_size
                and para_info["font_size"] != "Default"
            ):
                corrections.append(
                    {
                        "before": para_info["font_size"],
                        "after": expected_font_size,
                        "text": paragraph.text,
                        "issue": f"Font size was {para_info['font_size']} (expected: {expected_font_size})",
                    }
                )
                for run in paragraph.runs:
                    run.font.size = Pt(expected_font_size)
        except Exception as e:
            raise RuntimeError(f"Error correcting font size: {e}")

        return corrections

    def correct_font_color(self, paragraph):
        corrections = []
        try:
            para_info = self.get_paragraph_font_info(paragraph)
            expected_font_colors = self.style_guide["paragraph"]["font_colors"]

            if (
                para_info["font_color"] not in expected_font_colors
                and para_info["font_color"] != "Default"
            ):
                corrections.append(
                    {
                        "before": para_info["font_color"],
                        "after": expected_font_colors[0],
                        "text": paragraph.text,
                        "issue": f"Font color was {para_info['font_color']} (expected: one of {expected_font_colors})",
                    }
                )
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(expected_font_colors[0])
        except Exception as e:
            raise RuntimeError(f"Error correcting font color: {e}")

        return corrections

    def correct_font_family(self, paragraph):
        corrections = []
        try:
            para_info = self.get_paragraph_font_info(paragraph)
            expected_font_names = self.style_guide["paragraph"]["font_name"]

            if para_info["font_name"] not in expected_font_names:
                corrections.append(
                    {
                        "before": para_info["font_name"],
                        "after": expected_font_names[0],
                        "text": paragraph.text,
                        "issue": f"Font family was {para_info['font_name']} (expected: {expected_font_names})",
                    }
                )
                for run in paragraph.runs:
                    run.font.name = expected_font_names[0]
        except Exception as e:
            raise RuntimeError(f"Error correcting font family: {e}")

        return corrections

    def correct_numeral_symbol_spacing(self, paragraph):
        corrections = []
        try:
            text = paragraph.text
            new_text = re.sub(r"(\d+)\s+([%$])", r"\1\2", text)
            if new_text != text:
                corrections.append(
                    {
                        "before": text,
                        "after": new_text,
                        "text": paragraph.text,
                        "issue": f"Incorrect spacing between numeral and symbol.",
                    }
                )
                paragraph.clear()  # Remove old text
                paragraph.add_run(new_text)  # Add corrected text
        except Exception as e:
            raise RuntimeError(f"Error correcting numeral-symbol spacing: {e}")

        return corrections

    def apply_rolling_corrections(self, paragraph):
        """
        Apply all corrections one by one in a rolling manner.
        Each step passes the corrected paragraph to the next correction function.
        """
        corrections_log = []

        # Detect proper nouns to avoid grammar fixes for those
        proper_nouns = self.detect_proper_nouns(paragraph)

        # Rolling updates for the paragraph
        corrections_log += self.correct_font_size(paragraph)
        corrections_log += self.correct_font_color(paragraph)
        corrections_log += self.correct_font_family(paragraph)
        corrections_log += self.correct_numeral_symbol_spacing(paragraph)
        corrections_log += self.detect_abbreviations(paragraph)
        corrections_log += self.format_emails(paragraph)
        corrections_log += self.correct_spelling_and_grammar(paragraph, proper_nouns)

        return corrections_log

    def pre_process_document(self, input_doc):
        try:
            # Create a deep copy of the input document
            modified_doc = deepcopy(input_doc)
            log = []
    
            # Process regular paragraphs in the document body
            for para in modified_doc.paragraphs:
                if not para.text.strip():
                    continue
    
                # Apply rolling corrections to the paragraph
                corrections = self.apply_rolling_corrections(para)
    
                if corrections:
                    log.append(
                        {"paragraph_text": para.text, "corrections": corrections}
                    )
    
            # Process paragraphs inside tables
            for table in modified_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if not para.text.strip():
                                continue
    
                            # Apply rolling corrections to each paragraph inside the table
                            corrections = self.apply_rolling_corrections(para)
    
                            if corrections:
                                log.append(
                                    {"paragraph_text": para.text, "corrections": corrections}
                                )
    
            # Generate the correction log document
            log_doc = self.generate_corrections_doc(log)
            return modified_doc, log_doc
        except Exception as e:
            raise RuntimeError(f"Error during document preprocessing: {e}")


    def generate_corrections_doc(self, corrections_log):
        try:
            # Create a new Word document for logs
            doc = Document()

            # Add a title to the document
            doc.add_heading('Document Correction Log', level=1)

            for entry in corrections_log:
                # Add the paragraph text as a heading
                doc.add_heading(f"Paragraph: {entry['paragraph_text']}", level=2)

                for correction in entry["corrections"]:
                    # Add the issue, before, and after values for each correction
                    doc.add_paragraph(f"Issue: {correction['issue']}")
                    doc.add_paragraph(f"Before: {correction['before']}")
                    doc.add_paragraph(f"After: {correction['after']}")
                    # Add a horizontal line to separate corrections
                    doc.add_paragraph('-' * 40)

            # Return the Document object for further processing
            return doc
        except Exception as e:
            raise RuntimeError(f"Error generating corrections document: {e}")

    def save_documents(self, modified_doc, log_doc, original_file_path):
        """
        Save the modified document and the correction log.
        Append '_corrected' and '_log' to the original file name.
        """
        corrected_file_path = original_file_path.replace(".docx", "_corrected.docx")
        log_file_path = original_file_path.replace(".docx", "_log.docx")

        # Save the corrected document
        modified_doc.save(corrected_file_path)

        # Save the log document
        log_doc.save(log_file_path)

        print(f"Corrected document saved at: {corrected_file_path}")
        print(f"Correction log saved at: {log_file_path}")
