# imports and set up the local model
from docx import Document
from llama_cpp import Llama, ChatCompletionRequestMessage
from copy import deepcopy
from smartdoc_utils import process_llm_output
from deterministic_preprocessor import DeterministicPreprocessor
from config import settings
from prompts import (
    karen_prompt,
    karen_system_prompt,
    llama_prompt,
    llama_system_prompt,
    gemma_prompt,
)
from llm_configs import karen_config, llama_config, gemma_config

# Initialize the model with GPU support
# llm = Llama(
#     model_path=settings.llm_model.model_name,
#     n_gpu_layers=-1,  # -1 means use all available GPU layers
#     n_ctx=16384,  # adjust based on your GPU memory
# )

# karen 7b q6_k creative
# llm = Llama.from_pretrained(
#     repo_id="FPHam/Karen_TheEditor_V2_CREATIVE_Mistral_7B-Q6_K-GGUF",
#     filename="karen_theeditor_v2_creative_mistral_7b.Q6_K.gguf",
#     n_gpu_layers=-1,
#     n_ctx=16384,
# )

# karen 7b q6_k strict
# llm = Llama.from_pretrained(
#     repo_id="TheBloke/Karen_TheEditor_V2_STRICT_Mistral_7B-GGUF",
#     filename="karen_theeditor_v2_strict_mistral_7b.Q6_K.gguf",
#     n_gpu_layers=-1,
#     n_ctx=16384,
# )

# llama 3.1 8b q6_k_l
# llm = Llama.from_pretrained(
#     repo_id="bartowski/Meta-Llama-3.1-8B-Instruct-GGUF",
#     filename="Meta-Llama-3.1-8B-Instruct-Q6_K_L.gguf",
#     n_gpu_layers=-1,
#     n_ctx=16384,
# )

# llama 3.2 3b fp16
# llm = Llama.from_pretrained(
#     repo_id="bartowski/Llama-3.2-3B-Instruct-GGUF",
#     filename="Llama-3.2-3B-Instruct-f16.gguf",
#     n_gpu_layers=-1,
#     n_ctx=16384,
# )

# gemma 2 9 b q6_k_l
# llm = Llama.from_pretrained(
#     repo_id="bartowski/gemma-2-9b-it-GGUF",
#     filename="gemma-2-9b-it-Q6_K_L.gguf",
#     n_gpu_layers=-1,
#     n_ctx=8192,
# )

# qwen2.5 7b q6_k_l
llm = Llama.from_pretrained(
    repo_id="bartowski/Qwen2.5-7B-Instruct-GGUF",
    filename="Qwen2.5-7B-Instruct-Q6_K_L.gguf",
    n_gpu_layers=-1,
    n_ctx=16384,
)

if "karen" in llm.model_path:
    generation_params = karen_config
    PROMPT_TEMPLATE = karen_prompt
    SYS_PROMPT = karen_system_prompt
elif "gemma" in llm.model_path:
    generation_params = gemma_config
    PROMPT_TEMPLATE = gemma_prompt
    SYS_PROMPT = None
else:
    generation_params = llama_config
    PROMPT_TEMPLATE = llama_prompt
    SYS_PROMPT = llama_system_prompt


def fetch_llm_response(
    text,
    prompt_template=PROMPT_TEMPLATE,
    sys_prompt=SYS_PROMPT,
):
    if sys_prompt is None:
        messages = [
            {
                "role": "user",
                "content": prompt_template.format(text=text),
            },
        ]
    else:
        messages = [
            {
                "role": "system",
                "content": sys_prompt,
            },
            {
                "role": "user",
                "content": prompt_template.format(text=text),
            },
        ]
    edited_text = llm.create_chat_completion(
        messages=messages,
        **generation_params,
    )
    response_text = edited_text["choices"][0]["message"]["content"].strip()
    print(f"PROMPT: \n {prompt_template.format(text=text)} \n\n")
    print(f"LLM RESPONSE: \n{response_text}\n\n")

    return response_text


# code that Aamir showed with language_tool checks comes here
# Pre-process an input document to handle formatting,
# language checks, etc. before further processing
#
# Args:
#   input_doc: The original input document
#   corrections_doc: An empty document to populate with corrections
#
# Returns:
#   modified_doc: A modified version of the input document
#   corrections_doc: Populated with any corrections made during p
def pre_process_document(input_doc, corrections_doc):

    # logic for pre-processing Fonts, Australian Language Checks, etc.

    # -- Need to uncomment following 3 lines - the code has not been tested here yet!

    preprocessor = DeterministicPreprocessor()
    modified_doc, corrections_doc = preprocessor.pre_process_document(input_doc)

    return modified_doc, corrections_doc


def process_document_paragraphs(modified_doc, corrections_doc):

    # modified_doc = deepcopy(input_doc)
    # Process each paragraph
    for para in modified_doc.paragraphs:
        # Store the original formatting
        original_runs = para.runs.copy()

        if not para.text.strip():
            print("Skipping LLM CALL:")
            continue
        else:
            # Get the text content
            text = para.text

        llm_output_text = fetch_llm_response(text)

        # edits, corrections = document_postprocessing(llm_output_text)
        edits, corrections = process_llm_output(llm_output_text)
        print("EDITS : \n", edits)
        print("CORRECTIONS: \n", corrections)
        # Clear the paragraph and add the edited text
        para.clear()
        # para.add_run(edited_text['choices'][0]['text'].strip())
        para.add_run(edits)
        # Attempt to reapply formatting
        new_runs = para.runs
        for i, run in enumerate(new_runs):
            if i < len(original_runs):
                run.font.name = original_runs[i].font.name
                run.font.size = original_runs[i].font.size
                run.font.bold = original_runs[i].font.bold
                run.font.italic = original_runs[i].font.italic
                run.font.color.rgb = original_runs[i].font.color.rgb
                run.font.underline = original_runs[i].font.underline
                # Add more formatting attributes as needed

        # Let us log the corrections made
        corrections_doc.add_paragraph()
        corrections_doc.add_paragraph(f"Original Text : \n {text}")
        corrections_doc.add_paragraph(f"Edits : \n {edits}")
        corrections_doc.add_paragraph(f"Corrections: \n{llm_output_text}")

    return modified_doc, corrections_doc


def process_document_tables(modified_doc, corrections_doc):

    ## REPLACE this with logic of the modified doc
    # modified_doc = deepcopy(input_doc)
    # Iterate through all tables in the document
    for table in modified_doc.tables:
        print("IN Table")
        printed_cells = set()  # To keep track of cells that have been processed
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                cell_id = (r_index, c_index)  # Unique identifier for the cell

                # Skip this cell if it is already processed as part of a merged cell
                if cell_id in printed_cells:
                    continue

                # Detect merged cells
                is_merged = False
                for other_cell in row.cells:
                    if other_cell is not cell and other_cell._element is cell._element:
                        is_merged = True
                        break

                # If it's a merged cell, avoid processing duplicates
                if is_merged:
                    # Register this cell's element to skip duplicates
                    for merged_row_index, merged_row in enumerate(table.rows):
                        for merged_cell_index, merged_cell in enumerate(
                            merged_row.cells
                        ):
                            if merged_cell._element is cell._element:
                                printed_cells.add((merged_row_index, merged_cell_index))

                # Append '**' to the text of the cell if not already processed
                if cell.text.strip():  # Check if the cell is not empty
                    #                    cell.text += '*T*B*L'
                    for para in cell.paragraphs:
                        # Add an asterisk (*) to the end of each cell paragraph
                        print(para.text)
                        # Just a small check to see that we processed this
                        # para.add_run('*T')

                        # Store the original formatting
                        original_runs = para.runs.copy()
                        # let us call the llm
                        llm_output_text = fetch_llm_response(para.text)

                        # edits, corrections = document_postprocessing(llm_output_text)
                        edits, corrections = process_llm_output(llm_output_text)

                        # let us reapply formatting

                        # Clear the paragraph and add the edited text
                        para.clear()
                        # para.add_run(edited_text['choices'][0]['text'].strip())
                        para.add_run(edits)
                        # Attempt to reapply formatting
                        new_runs = para.runs
                        for i, run in enumerate(new_runs):
                            if i < len(original_runs):
                                run.font.name = original_runs[i].font.name
                                run.font.size = original_runs[i].font.size
                                run.font.bold = original_runs[i].font.bold
                                run.font.italic = original_runs[i].font.italic
                                run.font.color.rgb = original_runs[i].font.color.rgb
                                run.font.underline = original_runs[i].font.underline
                                # Add more formatting attributes as needed

                        # Let us log the corrections made
                        corrections_doc.add_paragraph()
                        corrections_doc.add_paragraph(f"Original Text : \n {para.text}")
                        corrections_doc.add_paragraph(
                            f"Corrections: \n{llm_output_text}"
                        )

            print()  # Newline after each row

    return modified_doc, corrections_doc


def proofread_and_correct_document(modified_doc, corrections_doc):
    corrections_doc.add_heading("Corrections Made", 0)

    # Correct using deterministic grammar checking
    modified_doc, corrections_doc = pre_process_document(modified_doc, corrections_doc)
    modified_doc, corrections_doc = process_document_paragraphs(
        modified_doc, corrections_doc
    )
    modified_doc, corrections_doc = process_document_tables(
        modified_doc, corrections_doc
    )

    return modified_doc, corrections_doc


# Uncomment the below when testing this file for llm output
if __name__ == "__main__":
    input_path = "/home/cdsw/data/simple-word-file-with-table.docx"
    edit_path = "/home/cdsw/data/output_doc.docx"
    correction_path = "/home/cdsw/data/correction_file.docx"
    input_doc = Document(input_path)
    modified_doc = deepcopy(input_doc)
    # Create a document object for corrections
    corrections_doc = Document()
    modified_doc, corrections_doc = proofread_and_correct_document(
        modified_doc, corrections_doc
    )
    modified_doc.save(edit_path)
    corrections_doc.save(correction_path)
